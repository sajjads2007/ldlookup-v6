# LD Lookup — Version 6.0.1 (Lookup & Audit)
# Tabs:
#   1) Item lookup: single short text field + Find + Clear
#   2) Audit: upload files/images, take a photo, paste/edit list; auto-fix 7-digit to Lxxxxxxx; highlight corrections
# Output:
#   - Table: LNumber + Image (clickable 100px thumb → in-app modal)
#   - No "Name" column, no file exports
# Notes:
#   - Photo/PDF OCR is attempted if pytesseract is available; otherwise app warns and continues.
#   - Text PDFs are parsed with pdfplumber; scanned PDFs fall back to OCR if available.

import io
import re
from pathlib import Path
from typing import List, Tuple
from urllib.parse import quote_plus, unquote_plus

import pandas as pd
import requests
import streamlit as st

# Optional readers
import pdfplumber
from docx import Document

# Optional OCR (graceful fallback if not present)
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False
    Image = None  # type: ignore

# -------------------------------------------------
# Page & styles
# -------------------------------------------------
OXFORD = "#0B132B"
VERSION = "6.0.1"

st.set_page_config(page_title=f"LD Lookup v{VERSION}", page_icon="🧾", layout="wide")

st.markdown(
    f"""
    <style>
      .stApp, .main {{ background-color: {OXFORD}; }}
      .block-container {{
          background: #ffffff;
          border-radius: 16px;
          padding: 18px 18px 10px 18px;
          box-shadow: 0 2px 12px rgba(0,0,0,.12);
      }}
      .topbar {{
          background: {OXFORD};
          color: #fff;
          padding: 14px 18px;
          border-radius: 12px;
          margin-bottom: 10px;
      }}
      .topbar h1 {{ margin: 0; font-size: 22px; color: #fff; }}
      .disclaimer {{ color:#C9CED8; font-size:.92rem; margin:4px 0 12px 0; }}
      .tbl table {{ border-collapse: collapse; width: 100%; }}
      .tbl th, .tbl td {{ border: 1px solid #e6e6e6; padding: 8px; vertical-align: middle; font-size: .92rem; }}
      .tbl th {{ background: #f6f7fb; color: {OXFORD}; text-align: left; }}
      .thumb {{
          height: 100px; width: auto; object-fit: contain; border-radius: 6px;
          border: 1px solid #e6e6e6; background: #fff;
      }}
      .chip {{ display:inline-block; padding:4px 8px; margin:2px; border-radius:10px; background:#F2F7FF; border:1px solid #DFE7F5; }}
      .chip.fixed {{ background:#FFF7CC; border-color:#F3D25A; }}
      .muted {{ color:#6b7280; font-size:.9rem; }}
      @media (max-width: 640px) {{
        .tbl th, .tbl td {{ font-size: .88rem; }}
        .thumb {{ height: 96px; }}
      }}
      /* Hide default drag-hint text from uploader (visual only) */
      div[data-testid="stFileUploadDropzone"] small {{ display:none; }}
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    f"""
    <div class="topbar">
      <h1>LD Lookup — Version {VERSION} (Lookup & Audit)</h1>
      <div class="disclaimer">This is not an official app — currently in debugging mode</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# -------------------------------------------------
# Constants & helpers
# -------------------------------------------------
IMG_WIDTH = 640
IMG_QUALITY = 75
CDN_TEMPLATE = "https://cdn-tp2.mozu.com/28945-m4/cms/files/{L}.jpg?w={w}&q={q}"

# Match canonical L-numbers and 7-digit bare numbers
RE_LNUM = re.compile(r"\bL\d+\b", flags=re.IGNORECASE)
RE_BARE7 = re.compile(r"\b(\d{7})\b")

def build_image_url(l_number: str, width=IMG_WIDTH, quality=IMG_QUALITY) -> str:
    l = str(l_number).strip().upper()
    return CDN_TEMPLATE.format(L=l, w=width, q=quality)

def extract_lnumbers_from_text_with_correction(text: str) -> Tuple[List[str], List[str]]:
    """
    Returns (lnumbers, corrected) where:
      - lnumbers: deduped, order-preserving list of 'L\\d+'
      - corrected: subset that came from bare 7-digit numbers auto-prefixed with 'L'
    """
    if not text:
        return [], []
    found: List[str] = []
    corrected: List[str] = []

    # First collect canonical Lxxxx
    for m in RE_LNUM.finditer(text):
        found.append(m.group(0).upper())

    # Then collect bare 7-digit and prefix with L
    for m in RE_BARE7.finditer(text):
        candidate = f"L{m.group(1)}"
        found.append(candidate.upper())
        corrected.append(candidate.upper())

    # Deduplicate preserving order
    out, seen = [], set()
    corr_out: List[str] = []
    corr_set = set(corrected)
    for x in found:
        if x not in seen:
            seen.add(x)
            out.append(x)
            if x in corr_set:
                corr_out.append(x)
    return out, corr_out

def extract_lnumbers_from_dataframe(df: pd.DataFrame) -> Tuple[List[str], List[str]]:
    all_text = "\n".join(df.astype(str).fillna("").agg(" ".join, axis=1).tolist())
    return extract_lnumbers_from_text_with_correction(all_text)

def read_text_from_file(file) -> str:
    suffix = Path(file.name).suffix.lower()
    try:
        if suffix == ".csv":
            df = pd.read_csv(file, dtype=str, on_bad_lines="skip")
            return "\n".join(df.astype(str).fillna("").agg(" ".join, axis=1).tolist())
        if suffix in [".xlsx", ".xls"]:
            df = pd.read_excel(file, dtype=str)
            return "\n".join(df.astype(str).fillna("").agg(" ".join, axis=1).tolist())
        if suffix == ".txt":
            return file.read().decode("utf-8", errors="ignore")
        if suffix == ".docx":
            doc = Document(file)
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if suffix == ".pdf":
            # Text-based first
            with pdfplumber.open(file) as pdf:
                text = "\n".join((p.extract_text() or "") for p in pdf.pages)
            return text
        if suffix in [".png", ".jpg", ".jpeg", ".webp"]:
            if not OCR_AVAILABLE:
                return ""
            image = Image.open(file)
            return pytesseract.image_to_string(image)
        return ""
    except Exception:
        return ""

def ocr_bytes_to_text(b: bytes) -> str:
    if not OCR_AVAILABLE or Image is None:
        return ""
    try:
        img = Image.open(io.BytesIO(b))
        return pytesseract.image_to_string(img)
    except Exception:
        return ""

def render_table(l_numbers: List[str]) -> None:
    # Build clickable thumbnail table via HTML (fast, flexible)
    rows = []
    for l in l_numbers:
        url = build_image_url(l)
        href = f"?preview={quote_plus(url)}"
        img_cell = f'<a href="{href}" title="Preview {l}"><img class="thumb" src="{url}" alt="{l}"></a>'
        rows.append(f"<tr><td>{l}</td><td>{img_cell}</td></tr>")
    st.markdown(
        f"""
        <div class="tbl">
          <table>
            <thead><tr><th>LNumber</th><th>Image</th></tr></thead>
            <tbody>{''.join(rows)}</tbody>
          </table>
        </div>
        """,
        unsafe_allow_html=True,
    )

def open_modal_if_requested():
    qp = st.query_params
    if "preview" in qp:
        preview_url = unquote_plus(qp.get("preview"))
        with st.modal("Preview", key="img_modal"):
            st.image(preview_url, use_column_width=True)
            if st.button("Close"):
                st.query_params.clear()
                st.rerun()

# -------------------------------------------------
# Tabs
# -------------------------------------------------
tab_lookup, tab_audit = st.tabs(["Item lookup", "Audit"])

# -------------------- Tab 1: Item lookup --------------------
with tab_lookup:
    st.caption("Enter a single code. If you type 7 digits (e.g., 1304179), I’ll treat it as L1304179.")
    col1, col2, col3 = st.columns([2,1,1])
    with col1:
        user_input = st.text_input("L-number", max_chars=16, placeholder="L1304179 or 1304179")
    with col2:
        find_clicked = st.button("Find", type="primary")
    with col3:
        clear_clicked = st.button("Clear")

    if clear_clicked:
        st.experimental_rerun()

    lnums: List[str] = []
    fixed: List[str] = []
    if find_clicked and user_input.strip():
        found, corrected = extract_lnumbers_from_text_with_correction(user_input.strip())
        if found:
            lnums = [found[0]]  # single check
            fixed = corrected
        else:
            st.warning("No valid L-number found.")

    if lnums:
        if lnums[0] in fixed:
            st.markdown(f"<div class='muted'>Corrected → <span class='chip fixed'>{lnums[0]}</span></div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='muted'>Detected → <span class='chip'>{lnums[0]}</span></div>", unsafe_allow_html=True)
        open_modal_if_requested()
        render_table(lnums)

# -------------------- Tab 2: Audit --------------------
with tab_audit:
    st.caption("Upload files or photos, take a picture, or paste a list. I’ll normalize 7-digit numbers to Lxxxxxxx. Edit before running.")
    c1, c2 = st.columns([1.1, 1])

    with c1:
        uploaded_files = st.file_uploader(
            "Upload CSV / XLSX / XLS / TXT / PDF / DOCX / PNG / JPG / JPEG / WEBP",
            type=["csv", "xlsx", "xls", "txt", "pdf", "docx", "png", "jpg", "jpeg", "webp"],
            accept_multiple_files=True,
            key="audit_file_uploader"
        )
        camera_img = st.camera_input("Take a picture of a list (optional)")
        if not OCR_AVAILABLE:
            st.markdown("<div class='muted'>Note: OCR not available on this server. Image text won’t be auto-read.</div>", unsafe_allow_html=True)

    with c2:
        pasted = st.text_area("Paste or edit L-numbers (free text, any separators)", height=220, key="audit_paste_area")

    raw_text_parts: List[str] = []
    if pasted.strip():
        raw_text_parts.append(pasted)

    if uploaded_files:
        for f in uploaded_files:
            txt = read_text_from_file(f)
            if txt:
                raw_text_parts.append(txt)

    if camera_img is not None:
        txt = ocr_bytes_to_text(camera_img.getvalue())
        if txt:
            raw_text_parts.append(txt)

    all_text = "\n".join(raw_text_parts)
    lnums_bulk: List[str] = []
    fixed_bulk: List[str] = []

    if all_text.strip():
        extracted, corrected = extract_lnumbers_from_text_with_correction(all_text)
        lnums_bulk = extracted
        fixed_bulk = corrected

    st.markdown("**Detected list (editable):**")
    editable_default = " ".join(lnums_bulk)
    edited_text = st.text_area("Edit the list below and click Run Audit", value=editable_default, height=120, key="audit_editable_list")

    if fixed_bulk:
        chips = "".join([f"<span class='chip fixed'>{x}</span>" for x in fixed_bulk])
        st.markdown(f"<div class='muted'>Auto-corrected from 7 digits → {chips}</div>", unsafe_allow_html=True)

    run_audit = st.button("Run Audit", type="primary")
    if run_audit:
        final_list, _ = extract_lnumbers_from_text_with_correction(edited_text)
        if not final_list:
            st.warning("No L-numbers to process after edits.")
        else:
            open_modal_if_requested()
            render_table(final_list)
